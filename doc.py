import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests

# ================== Gemini API Helper ==================
def generate_text_gemini(prompt, api_key, max_tokens=1000):
    """
    Generate text from Gemini API based on prompt
    """
    url = "https://api.gemini.com/v1/generate"  # Replace with actual endpoint
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "prompt": prompt,
        "max_tokens": max_tokens
    }
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        result = response.json()
        return result.get("text", "")
    else:
        print("Error:", response.status_code, response.text)
        return ""

# ================== Predefined Chapter Titles ==================
chapter_titles = [
    "Introduction",
    "System Analysis",
    "System Design",
    "Implementation",
    "Testing and Results",
    "Conclusions and Future Work"
]

# ================== CLI Input ==================
print("=== Insurance Management System Documentation Generator ===\n")

project_title = input("Enter Project Title: ")
author_name = input("Enter Author Name: ")
college_name = input("Enter College Name: ")
department_name = input("Enter Department Name: ")
year = input("Enter Year: ")

api_key = input("Enter your Gemini API Key: ")
output_file = input("Enter output Word file name (e.g., insurance_doc.docx): ")

# Ask number of pages per chapter
chapter_pages = []
print("\nEnter expected number of pages for each chapter:")
for title in chapter_titles:
    pages = int(input(f"{title}: "))
    chapter_pages.append(pages)

# ================== Create Word Document ==================
doc = Document()

# Title Page
doc.add_paragraph(project_title, style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"Author: {author_name}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"College: {college_name}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"Department: {department_name}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"Year: {year}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Acknowledgements
doc.add_heading("Acknowledgements", level=1)
doc.add_paragraph("We would like to express our sincere gratitude to our guide, mentors, and all who supported us in completing this project.")
doc.add_page_break()

# Abstract
doc.add_heading("Abstract", level=1)
doc.add_paragraph("This project focuses on designing and developing an efficient insurance management system that automates policy management, claim processing, and reporting for improved efficiency and accuracy.")
doc.add_page_break()

# Generate chapters
for idx, title in enumerate(chapter_titles):
    doc.add_heading(title, level=1)
    prompt = f"Write a detailed chapter on '{title}' for an insurance management system documentation. The chapter should be approximately {chapter_pages[idx]} pages long in Word format. Include explanations, diagrams placeholders, and examples."
    content = generate_text_gemini(prompt, api_key, max_tokens=1500)
    doc.add_paragraph(content)
    doc.add_page_break()

# References
doc.add_heading("References", level=1)
doc.add_paragraph("1. Book/Journal/Website 1\n2. Book/Journal/Website 2\n3. ...")

# Save document
doc.save(output_file)
print(f"\n✅ Documentation generated and saved as {output_file}")
